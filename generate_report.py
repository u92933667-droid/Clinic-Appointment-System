from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def make_report():
    doc = Document()

    # --- Global Styling ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("🏥 CLINIC APPOINTMENT SYSTEM\n")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Complex Engineering Problem (CEP) - Final Project Report")
    run.italic = True
    run.font.size = Pt(16)

    doc.add_paragraph("\n" * 3)

    # Info Table on Cover
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = [
        ["Course:", "Computer Programming in Python (CEP)"],
        ["Instructor:", "Dr. M. Abbas Abbasi"],
        ["Semester:", "Fall 2025"],
        ["Submission Date:", "January 1, 2026"],
        ["Team Members:", "Student 1: [Your Name] - [Reg #]\nStudent 2: [Partner Name] - [Reg #]"]
    ]
    for i, row in enumerate(data):
        table.cell(i, 0).text = row[0]
        table.cell(i, 1).text = row[1]
        table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # --- EXECUTIVE SUMMARY ---
    add_heading(doc, 'EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph(
        "This report presents a comprehensive web-based Clinic Appointment Management System developed as a "
        "Complex Engineering Problem (CEP) project. The system provides an end-to-end solution for managing "
        "patient appointments in a clinical setting, featuring a modern user interface with full CRUD operations."
    )
    
    summary_bullets = [
        "Fully functional appointment booking system",
        "8 medical specialties with visual identification",
        "Modern dark-theme user interface",
        "RESTful API with comprehensive documentation"
    ]
    for bullet in summary_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # --- 1. INTRODUCTION ---
    add_heading(doc, '1. INTRODUCTION', level=1)
    add_heading(doc, '1.1 Problem Statement', level=2)
    doc.add_paragraph(
        "Healthcare facilities require efficient appointment management systems to streamline patient scheduling. "
        "Traditional paper-based systems are prone to errors and poor record-keeping."
    )

    # --- 5. DATABASE DESIGN ---
    add_heading(doc, '5. DATABASE DESIGN', level=1)
    add_heading(doc, '5.1 Table Schemas', level=2)
    
    db_table = doc.add_table(rows=1, cols=2)
    db_table.style = 'Light Shading Accent 1'
    hdr_cells = db_table.rows[0].cells
    hdr_cells[0].text = 'Entity'
    hdr_cells[1].text = 'Fields'
    
    entities = [
        ["Patients", "ID, Name, Age, Phone"],
        ["Doctors", "ID, Name, Specialty"],
        ["Appointments", "ID, Patient_ID, Doctor_ID, Date, Time, Reason"]
    ]
    for ent, fields in entities:
        row_cells = db_table.add_row().cells
        row_cells[0].text = ent
        row_cells[1].text = fields

    # --- 11. CLO MAPPING TABLE ---
    add_heading(doc, '11. CLO MAPPING', level=1)
    clo_table = doc.add_table(rows=1, cols=3)
    clo_table.style = 'Medium Grid 1 Accent 1'
    hdr = clo_table.rows[0].cells
    hdr[0].text = 'CLO'
    hdr[1].text = 'Description'
    hdr[2].text = 'Evidence'

    clos = [
        ["CLO-1", "Python Concepts (OOP, File I/O)", "FastAPI backend using classes and SQLite storage."],
        ["CLO-2", "Emerging Tech", "Used FastAPI (Async) and React with Vite."],
        ["CLO-3", "Development & Debugging", "Fixed CORS issues and N+1 query problems."],
        ["CLO-4", "Documentation", "README, code comments, and this technical report."],
        ["CLO-5", "Project Planning", "6-week cycle following Database-API-Frontend workflow."]
    ]
    for c_id, desc, ev in clos:
        row = clo_table.add_row().cells
        row[0].text = c_id
        row[1].text = desc
        row[2].text = ev

    # --- 12. CHALLENGES ---
    add_heading(doc, '12. CHALLENGES AND SOLUTIONS', level=1)
    p = doc.add_paragraph()
    run = p.add_run("Challenge: CORS Errors\n")
    run.bold = True
    doc.add_paragraph(
        "Solution: Implemented CORSMiddleware in FastAPI to allow requests from the React frontend origin."
    )

    # --- CONCLUSION ---
    add_heading(doc, '14. CONCLUSION', level=1)
    doc.add_paragraph(
        "The Clinic Appointment System successfully meets all objectives and demonstrates a comprehensive "
        "understanding of full-stack web development. The project statistics show a robust application "
        "with over 800 lines of code across Python, JavaScript, and CSS."
    )

    # Save the document
    file_name = "Clinic_Appointment_System_CEP_Report.docx"
    doc.save(file_name)
    print(f"Success! Report generated as {file_name}")

if __name__ == "__main__":
    make_report()